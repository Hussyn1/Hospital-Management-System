from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
from datetime import date


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "Hospital_Management_System_Project_Report.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color="DADCE0", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, width=9360, indent=120):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        el = margins.find(qn(f"w:{name}"))
        if el is None:
            el = OxmlElement(f"w:{name}")
            margins.append(el)
        el.set(qn("w:w"), str(value))
        el.set(qn("w:type"), "dxa")


def apply_table_style(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table)
    set_table_borders(table)
    set_cell_margins(table)
    for row_index, row in enumerate(table.rows):
        for i, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if i < len(widths):
                set_cell_width(cell, widths[i])
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.1
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9)
            if row_index == 0:
                set_cell_shading(cell, "E8EEF5")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(17, 24, 39)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            cells[i].text = str(value)
    apply_table_style(table, widths)
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def file_size(path):
    p = ROOT / path
    if not p.exists():
        return ""
    size = p.stat().st_size
    if size >= 1024 * 1024:
        return f"{size / (1024 * 1024):.1f} MB"
    if size >= 1024:
        return f"{size / 1024:.1f} KB"
    return f"{size} B"


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(11)
styles["Normal"].paragraph_format.space_after = Pt(6)
styles["Normal"].paragraph_format.line_spacing = 1.25

for name, size, color in (
    ("Heading 1", 16, RGBColor(46, 116, 181)),
    ("Heading 2", 13, RGBColor(46, 116, 181)),
    ("Heading 3", 12, RGBColor(31, 77, 120)),
):
    style = styles[name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = color
    style.paragraph_format.space_before = Pt(12)
    style.paragraph_format.space_after = Pt(6)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Hospital Management System OOP Project")
run.bold = True
run.font.name = "Calibri"
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(11, 37, 69)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.add_run("Comprehensive Technical Report").bold = True

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f"Generated on {date.today().isoformat()} | Project path: {ROOT}")

doc.add_paragraph(
    "This report explains how the JavaFX Hospital Management System works, what each folder and source file is responsible for, how the libraries support the app, and how the data moves from screens to services, DAOs, and SQLite."
)

doc.add_heading("1. Executive Summary", level=1)
doc.add_paragraph(
    "The project is a desktop hospital management application written in Java using JavaFX for the user interface and SQLite for local persistence. It follows a layered object-oriented design: FXML and controllers handle screens, services apply business rules, DAO classes run SQL, model classes carry data, and utility classes provide shared helpers."
)
add_bullets(doc, [
    "Main entry point: com.hospital.MainLauncher launches the JavaFX Application class MainApp.",
    "User interface: LoginView.fxml and DashboardView.fxml define the base screens; DashboardController dynamically builds most module screens in Java.",
    "Database: data/hms.db stores the local SQLite database, initialized and seeded by DBConnection.java.",
    "Core modules: patients, doctors, appointments, EMR, pharmacy, laboratory, wards and beds, emergency triage, billing, HR shifts, inventory, reports, and notifications.",
    "Native Eclipse support: .project, .classpath, .settings, lib jars, and HospitalManagementSystem.launch allow the app to run without Gradle.",
])

doc.add_heading("2. High-Level Architecture", level=1)
add_table(doc, ["Layer", "Main files", "Purpose"], [
    ("Application startup", "MainLauncher.java, MainApp.java", "Bootstraps JavaFX, initializes the database, loads LoginView.fxml, and applies theme.css."),
    ("View layer", "LoginView.fxml, DashboardView.fxml, theme.css", "Defines the visible layout, base dashboard shell, login form, styling, sidebar, cards, buttons, dialogs, tables, and inputs."),
    ("Controller layer", "LoginController.java, DashboardController.java", "Responds to button clicks, reads form fields, updates tables, opens dialogs, and calls service classes."),
    ("Service layer", "services/*.java", "Contains business rules such as validation, double-booking checks, stock checks, admission/discharge transactions, billing calculations, and workflow coordination."),
    ("DAO layer", "dao/*.java", "Performs direct SQL operations against SQLite and maps rows into model objects."),
    ("Model layer", "models/*.java, enums/*.java", "Plain Java objects and enums representing hospital domain data and controlled status values."),
    ("Utilities", "AlertHelper.java, HashHelper.java, Validator.java", "Shared alert dialogs, password hashing, and input validation."),
], [1700, 2600, 5060])

doc.add_paragraph("Typical runtime flow:")
add_numbered(doc, [
    "User starts MainLauncher from Eclipse or Gradle.",
    "MainLauncher calls Application.launch(MainApp.class, args).",
    "MainApp initializes DBConnection, which creates the SQLite schema and inserts mock data if the database is empty.",
    "MainApp loads LoginView.fxml and displays the login screen.",
    "LoginController authenticates the user through AuthService.",
    "On success, DashboardView.fxml is loaded and DashboardController builds role-appropriate navigation and module screens.",
    "User actions call services, services call DAOs, DAOs update SQLite, then the controller refreshes the UI.",
])

doc.add_heading("3. Folder and File Inventory", level=1)
add_table(doc, ["Folder / File", "Contains", "What it does"], [
    (".git/", "Git metadata", "Tracks project history. Not used at runtime."),
    (".gradle/", "Gradle cache files", "Generated by Gradle; can be ignored for Eclipse-native use."),
    (".settings/", "Eclipse Java compiler preferences", "Sets Java 17 compiler compliance for Eclipse."),
    ("bin/", "Compiled .class files and copied resources", "Eclipse/Javac output. JavaFX loads resources from here when running compiled classes."),
    ("build/", "Gradle build output", "Generated when using Gradle tasks."),
    ("data/", "hms.db", "SQLite database file used by the app."),
    ("lib/", "JavaFX, SQLite JDBC, SLF4J jars", "Self-contained dependencies for Eclipse classpath."),
    ("src/main/java/", "Java source packages", "Application code organized by controllers, services, DAOs, models, enums, and utilities."),
    ("src/main/resources/", "FXML and CSS", "UI layout files and the application stylesheet."),
    ("src/test/java/", "JUnit test", "Contains DBConnectionTest for database initialization checks."),
    ("build.gradle", "Gradle build script", "Defines Java, JavaFX, SQLite, JUnit, toolchain, and application main class for Gradle users."),
    ("settings.gradle", "Gradle settings", "Names the Gradle project HospitalManagementSystem."),
    (".project", "Eclipse project descriptor", "Allows Eclipse to import the folder as a Java project."),
    (".classpath", "Eclipse classpath descriptor", "Lists source folders, Java 17 runtime, JUnit 5, and local lib jars."),
    ("HospitalManagementSystem.launch", "Eclipse launch config", "Runs com.hospital.MainLauncher with the project as working directory."),
], [2300, 2500, 4560])

doc.add_heading("4. Libraries and Dependencies", level=1)
add_table(doc, ["Library / file", "Role in the project", "Why it is needed"], [
    ("Java 17", "Language/runtime", "Runs the application and provides modern Java APIs such as LocalDate and LocalDateTime."),
    ("JavaFX Controls 17.0.9", "UI controls", "Provides Button, TableView, ComboBox, TextField, Alert, DatePicker, and other widgets."),
    ("JavaFX FXML 17.0.9", "FXML loading", "Allows LoginView.fxml and DashboardView.fxml to be loaded and connected to controller classes."),
    ("JavaFX Graphics 17.0.9", "Rendering/windowing", "Provides Stage, Scene, Parent, Image, layout rendering, and graphical runtime support."),
    ("JavaFX Base 17.0.9", "Core JavaFX APIs", "Provides observable collections and foundational JavaFX classes used by controls and tables."),
    ("JavaFX Swing 17.0.9", "Swing bridge module", "Included by the Gradle JavaFX module list; useful if JavaFX needs Swing integration."),
    ("SQLite JDBC 3.43.0.0", "Database driver", "Lets Java code open jdbc:sqlite:data/hms.db and execute SQL statements."),
    ("SLF4J API 1.7.30", "Logging API dependency", "Required by SQLite JDBC in some configurations; prevents missing logging class errors."),
    ("JUnit Jupiter 5.9.3", "Testing framework", "Used by DBConnectionTest to verify schema creation and seed data."),
    ("Gradle Java plugin", "Build support", "Compiles Java source in Gradle mode."),
    ("Gradle application plugin", "Run support", "Defines MainLauncher as the runnable class for Gradle."),
    ("OpenJFX Gradle plugin", "JavaFX Gradle support", "Downloads and configures JavaFX modules for Gradle builds."),
], [2400, 2600, 4360])

doc.add_heading("5. Source Package Overview", level=1)
add_table(doc, ["Package", "Files", "Responsibility"], [
    ("com.hospital", "MainApp.java, MainLauncher.java", "Application startup and JavaFX bootstrapping."),
    ("com.hospital.controllers", "LoginController.java, DashboardController.java", "Screen behavior, navigation, forms, dialogs, table refresh, and user actions."),
    ("com.hospital.dao", "AdmissionDAO through WardDAO", "SQL persistence layer for each domain object plus DBConnection for schema setup."),
    ("com.hospital.enums", "AppointmentStatus, BedStatus, PaymentStatus, Role, TriageLevel, WardType", "Controlled values used by models, services, and database CHECK constraints."),
    ("com.hospital.models", "Admission through Ward", "Domain data objects with fields, getters, and setters."),
    ("com.hospital.services", "AppointmentService through WardService", "Business logic layer that validates input and coordinates DAO operations."),
    ("com.hospital.utils", "AlertHelper, HashHelper, Validator", "Common helpers for dialogs, password hashing, and validation."),
], [2300, 2600, 4460])

doc.add_heading("6. Application Startup Files", level=1)
add_table(doc, ["File", "What it does"], [
    ("src/main/java/com/hospital/MainLauncher.java", "A non-modular JavaFX launcher. It does not extend Application; it calls Application.launch(MainApp.class, args). This avoids common JavaFX module-launcher problems in plain classpath/Eclipse setups."),
    ("src/main/java/com/hospital/MainApp.java", "Extends javafx.application.Application. In init(), it opens DBConnection to force schema and mock-data initialization. In start(), it loads /fxml/LoginView.fxml, creates the Scene, applies /css/theme.css, sets the window title, and displays the primary Stage."),
], [3000, 6360])

doc.add_heading("7. Controller Files", level=1)
add_table(doc, ["File", "What it does"], [
    ("LoginController.java", "Controls the login screen. It reads username/password fields, calls AuthService.login, shows validation errors, and loads DashboardView.fxml after successful authentication."),
    ("DashboardController.java", "The main controller for the whole application after login. It creates sidebar buttons based on user role, switches the central content area, builds module screens programmatically, handles forms/dialogs, displays tables, calls services, and refreshes data after changes. It contains the UI logic for dashboard metrics, patient management, doctor management, appointments, EMR, pharmacy, lab, wards, emergency triage, billing, HR shifts, inventory, and reports."),
], [3000, 6360])

doc.add_heading("8. Service Layer Files", level=1)
add_table(doc, ["File", "Business responsibility"], [
    ("AuthService.java", "Authenticates users by username/password, hashes the entered password, compares it to the stored password hash, stores the current logged-in user, and supports logout."),
    ("PatientService.java", "Registers, updates, searches, retrieves, and soft-deletes patients. It validates required fields and delegates persistence to PatientDAO."),
    ("DoctorService.java", "Registers and updates doctors. It validates profile and account fields, creates linked user/doctor data through DoctorDAO, and supports searching by name, department, or specialty."),
    ("AppointmentService.java", "Books and updates appointments. It validates required appointment data and prevents doctor double-booking for the same date/time before saving."),
    ("EMRService.java", "Creates and retrieves medical records. It validates patient, doctor, diagnosis, and visit data before using EMRDAO."),
    ("PharmacyService.java", "Manages medicines and prescriptions. It validates medicine data, creates prescriptions, checks medicine stock, decrements stock when dispensing, and marks prescriptions filled."),
    ("LabService.java", "Creates lab test requests and records diagnostic results. It validates test names and result text."),
    ("WardService.java", "Coordinates bed admissions and discharges. It checks patient and bed status, updates bed occupancy, updates patient status, writes admissions, and generates discharge billing records."),
    ("EmergencyService.java", "Registers emergency triage queue entries, assigns an available bed when possible, sorts waiting cases by priority, and resolves emergency records."),
    ("BillingService.java", "Generates appointment/admission bills, processes payments, calculates payment status, and exposes revenue breakdown data."),
    ("HRService.java", "Assigns and removes staff shifts. It validates shift user, times, and role values."),
    ("InventoryService.java", "Manages hospital supplies, low-stock checks, stock adjustments, and purchase orders."),
    ("ReportsService.java", "Builds dashboard metrics and report aggregates such as appointments by status, bed occupancy, revenue, pending items, and unread notifications."),
], [2600, 6760])

doc.add_heading("9. DAO Layer Files", level=1)
add_table(doc, ["File", "Database responsibility"], [
    ("DBConnection.java", "Central database utility. Defines jdbc:sqlite:data/hms.db, creates the data folder, initializes schema tables, inserts mock data when users table is empty, returns new JDBC connections, and provides getLastInsertId for SQLite-safe row ID retrieval."),
    ("PatientDAO.java", "CRUD and search for patients. Generates patient IDs such as PAT001, saves patient rows, updates demographics/status, soft-deletes using is_deleted, and maps rows to Patient objects."),
    ("DoctorDAO.java", "Transactional doctor creation/update. Creates linked users and doctors records, supports search by name/specialty/department, and maps joined doctor/user rows."),
    ("AppointmentDAO.java", "Saves, updates, changes status, retrieves all appointments, retrieves patient appointments, and finds doctor appointments by date for conflict checks."),
    ("EMRDAO.java", "Saves medical records and loads patient histories with joined patient/doctor names."),
    ("PrescriptionDAO.java", "Transactional save for prescription header and line items, updates prescription status, retrieves prescriptions with items, and maps item details with medicine data."),
    ("MedicineDAO.java", "CRUD/search for medicine catalog, low-stock query, and stock adjustment."),
    ("LabDAO.java", "Creates lab requests, records result text/abnormal flags/report path, and loads pending/all lab requests with patient/doctor context."),
    ("WardDAO.java", "Saves and lists hospital wards."),
    ("BedDAO.java", "Saves beds, updates bed status, lists beds by ward/all beds, and counts available/occupied beds."),
    ("AdmissionDAO.java", "Saves admissions, records discharge time, retrieves active admissions, and maps joined patient/bed/ward details."),
    ("EmergencyDAO.java", "Saves emergency queue records, updates status/bed assignment, returns active queue sorted by triage priority and registration time."),
    ("BillingDAO.java", "Saves bills, applies payments, lists bills, and calculates revenue summary by payment status."),
    ("HRDAO.java", "Saves and deletes staff shift rows and lists shifts with staff names."),
    ("InventoryDAO.java", "Saves and updates inventory items, adjusts stock, lists low-stock items, and transactionally saves purchase orders while updating stock."),
    ("NotificationDAO.java", "Saves notifications, marks them read, lists notifications by user, and counts unread items."),
], [2600, 6760])

doc.add_heading("10. Model Files", level=1)
add_table(doc, ["Model file", "Represents"], [
    ("User.java", "Application user account: id, username, password hash, role, full name, contact, email, created time."),
    ("Patient.java", "Patient demographic and clinical profile: id, name, date of birth, blood group, contacts, history, status, deleted flag."),
    ("Doctor.java", "Doctor profile linked to a user: specialization, license number, consultation fee, availability schedule, department, contact fields."),
    ("Appointment.java", "Scheduled consultation: patient, doctor, date, time, status, notes, created time, joined names."),
    ("MedicalRecord.java", "EMR visit record: appointment, patient, doctor, diagnosis, visit date, treatment plan, joined names."),
    ("Medicine.java", "Pharmacy medicine item: name, generic name, stock, threshold, price, expiry date."),
    ("Prescription.java", "Prescription header: record, doctor, patient, status, created time, patient/doctor names, and items list."),
    ("PrescriptionItem.java", "Individual prescribed medicine: prescription id, medicine id, dosage, frequency, quantity, price helper fields."),
    ("LabRequest.java", "Diagnostic request/result: record, patient, test name, requested date, result text, abnormal flag, status, report file path."),
    ("Ward.java", "Hospital ward with id, name, and ward type."),
    ("Bed.java", "Bed record: ward, bed number, status, joined ward name/type."),
    ("Admission.java", "Hospital admission: patient, bed, admission/discharge dates, daily rate, joined patient/bed/ward names."),
    ("Bill.java", "Invoice/payment record: patient, optional appointment/admission, totals, paid amount, payment status/method, billing date."),
    ("EmergencyPatient.java", "Emergency queue entry: patient, triage level, priority, optional bed, registration time, status, joined patient/bed labels."),
    ("InventoryItem.java", "Non-pharmacy inventory item: item name/type, stock, unit, threshold, restock and expiry dates."),
    ("PurchaseOrder.java", "Inventory purchase order: item id/name, quantity, date, cost, supplier."),
    ("StaffShift.java", "Staff scheduling record: user, shift start/end, assigned role, staff name."),
    ("Notification.java", "User notification: recipient, message, read flag, created time."),
], [2600, 6760])

doc.add_heading("11. Enum Files", level=1)
add_table(doc, ["Enum", "Values / use"], [
    ("Role.java", "ADMIN, DOCTOR, NURSE, RECEPTIONIST, PHARMACIST, LAB_TECH. Drives permissions/navigation and user records."),
    ("AppointmentStatus.java", "PENDING, CONFIRMED, COMPLETED, CANCELLED. Used in appointments and UI actions."),
    ("BedStatus.java", "AVAILABLE, OCCUPIED, UNDER_MAINTENANCE. Used by bed management and admissions."),
    ("PaymentStatus.java", "UNPAID, PARTIAL, PAID. Used by billing/payment workflows."),
    ("TriageLevel.java", "CRITICAL, HIGH, MEDIUM, LOW. Each value has a numeric priority for ER queue sorting."),
    ("WardType.java", "GENERAL, ICU, PEDIATRIC, SURGICAL, MATERNITY. Used by ward records and database constraints."),
], [2600, 6760])

doc.add_heading("12. Utility Files", level=1)
add_table(doc, ["File", "What it does"], [
    ("AlertHelper.java", "Wraps JavaFX Alert dialogs for information, warning, error, and confirmation messages. It also applies theme.css to dialogs so popups match the app styling."),
    ("HashHelper.java", "Hashes passwords with SHA-256 and returns a lowercase hex string. Used when seeding users and checking login passwords."),
    ("Validator.java", "Provides reusable validation for email, phone/contact text, date of birth, non-empty strings, positive double values, and positive integer values."),
], [2600, 6760])

doc.add_heading("13. Resource Files", level=1)
add_table(doc, ["Resource", "What it does"], [
    ("src/main/resources/fxml/LoginView.fxml", "Defines the login screen layout: title, username field, password field, error label, login button, and default credential hint."),
    ("src/main/resources/fxml/DashboardView.fxml", "Defines the main dashboard shell: BorderPane, left sidebar, navigationBox for dynamic buttons, user card, sign-out button, and central content StackPane."),
    ("src/main/resources/css/theme.css", "Defines the app's light visual theme: white backgrounds, black text, blue active/primary buttons, gray inactive buttons, simple cards, inputs, tables, dialogs, scrollbars, tabs, and labels."),
], [3000, 6360])

doc.add_heading("14. Database Design", level=1)
doc.add_paragraph(
    "DBConnection.java creates eighteen tables. Most tables use primary keys, foreign keys, and CHECK constraints to keep status and category values consistent with Java enums."
)
add_table(doc, ["Table", "Purpose", "Key relationships / notes"], [
    ("users", "Login accounts and staff identity.", "Role is constrained to known roles; password_hash stores SHA-256 hash."),
    ("patients", "Patient demographic and clinical profile.", "Uses text primary key such as PAT001; supports soft delete via is_deleted."),
    ("doctors", "Doctor professional profile.", "user_id links to users and is unique."),
    ("appointments", "Consultation bookings.", "Links patient and doctor; status tracks booking lifecycle."),
    ("medical_records", "EMR visit records.", "Optionally links to appointment; links patient and doctor."),
    ("medicines", "Pharmacy catalog.", "Tracks stock quantity, threshold, price, and expiry."),
    ("prescriptions", "Prescription headers.", "Links medical record, doctor, and patient; status PENDING/FILLED."),
    ("prescription_items", "Prescription detail lines.", "Links prescription and medicine; cascade delete with prescription."),
    ("lab_requests", "Diagnostic orders and results.", "Links patient and optionally medical record; supports pending/completed status."),
    ("wards", "Hospital ward definitions.", "Ward type constrained to enum values."),
    ("beds", "Beds inside wards.", "Unique bed number per ward; status tracks availability."),
    ("admissions", "Inpatient stays.", "Links patient and bed; discharge_date null means active."),
    ("billing", "Invoices and payments.", "Links patient and optionally appointment/admission; tracks paid amount and status."),
    ("emergency_queue", "ER triage queue.", "Links patient and optional bed; sorted by queue_priority and registered_at."),
    ("inventory", "Hospital supplies/equipment.", "Tracks item type, stock, threshold, restock and expiry dates."),
    ("purchase_orders", "Inventory restocking records.", "Links to inventory item."),
    ("staff_shifts", "Staff scheduling.", "Links to users and stores shift start/end."),
    ("notifications", "User notifications.", "Links to user and tracks read/unread state."),
], [1800, 3300, 4260])

doc.add_heading("15. Main Workflows", level=1)
workflow_rows = [
    ("Login", "LoginController -> AuthService -> DBConnection/users", "User enters credentials, password is hashed, database user is loaded, and currentUser is set."),
    ("Patient registration", "DashboardController -> PatientService -> PatientDAO", "Form data becomes a Patient object; service validates it; DAO generates next patient ID and inserts the row."),
    ("Doctor registration", "DashboardController -> DoctorService -> DoctorDAO", "Creates a user account and matching doctor profile inside one transaction."),
    ("Appointment booking", "DashboardController -> AppointmentService -> AppointmentDAO", "Service checks required fields and rejects duplicate doctor date/time bookings before insert."),
    ("Consultation / EMR", "DashboardController -> EMRService, PharmacyService, LabService, BillingService", "Doctor records diagnosis/treatment, can add prescription and lab orders, completes appointment, and generates billing."),
    ("Prescription dispensing", "DashboardController -> PharmacyService -> PrescriptionDAO/MedicineDAO", "Service checks stock for all items, decrements medicine inventory, and marks prescription filled."),
    ("Lab result entry", "DashboardController -> LabService -> LabDAO", "Lab technician selects a pending request and records result text, abnormal flag, status, and optional report path."),
    ("Ward admission", "DashboardController -> WardService -> PatientDAO/BedDAO/AdmissionDAO", "Service checks patient/bed eligibility, marks bed occupied, marks patient admitted, and inserts admission."),
    ("Discharge", "DashboardController -> WardService -> BillingService", "Closes active admission, frees bed, marks patient discharged, and creates ward-stay invoice."),
    ("Emergency triage", "DashboardController -> EmergencyService -> EmergencyDAO/BedDAO", "Registers a triage case, calculates queue priority, assigns an available bed when possible, and lists active waiting cases by priority."),
    ("Billing payment", "DashboardController -> BillingService -> BillingDAO", "Payment amount is validated, new paid total determines UNPAID/PARTIAL/PAID status, and invoice is updated."),
    ("HR scheduling", "DashboardController -> HRService -> HRDAO", "Shift form creates StaffShift records after time/user validation."),
    ("Inventory restocking", "DashboardController -> InventoryService -> InventoryDAO", "Creates purchase orders and increments stock in one transaction."),
    ("Reports", "DashboardController -> ReportsService", "Reads aggregate counts and sums for dashboard cards and charts."),
]
add_table(doc, ["Workflow", "Code path", "What happens"], workflow_rows, [1900, 3100, 4360])

doc.add_heading("16. Role-Based Navigation", level=1)
doc.add_paragraph(
    "DashboardController builds the sidebar dynamically after login. The visible modules depend on AuthService.getCurrentUser().getRole(). Administrators see the broadest navigation; clinical and operational roles see the modules relevant to their job."
)
add_bullets(doc, [
    "ADMIN: full management areas such as dashboard, patients, doctors, appointments, wards, billing, HR, inventory, and reports.",
    "DOCTOR: appointment, EMR, prescription/lab ordering, and patient-history oriented screens.",
    "NURSE: patient, ward, admission, and emergency triage oriented screens.",
    "RECEPTIONIST: patient registration, appointment booking, and billing-oriented screens.",
    "PHARMACIST: medicine catalog, prescriptions, dispensing, and pharmacy stock alerts.",
    "LAB_TECH: lab request queue and diagnostic result entry.",
])

doc.add_heading("17. Configuration and Build Files", level=1)
add_table(doc, ["File", "Purpose"], [
    ("build.gradle", "Gradle build script. Applies java, application, and OpenJFX plugins; sets Java 17; configures JavaFX 17.0.9 modules; declares SQLite JDBC and JUnit dependencies; sets main class to com.hospital.MainLauncher."),
    ("settings.gradle", "Sets rootProject.name to HospitalManagementSystem."),
    (".project", "Eclipse metadata that identifies the project and enables the Java builder."),
    (".classpath", "Eclipse source/dependency map. Adds src/main/java, src/main/resources, src/test/java, JavaSE-17, JUnit 5, and lib/*.jar dependencies."),
    (".settings/org.eclipse.jdt.core.prefs", "Eclipse compiler settings for Java 17 compliance and release mode."),
    ("HospitalManagementSystem.launch", "Eclipse launch configuration that runs com.hospital.MainLauncher from the project working directory."),
], [3000, 6360])

doc.add_heading("18. File Reference Appendix", level=1)
all_files = [
    ("MainApp.java", "Starts the JavaFX app, initializes database, loads login UI, applies CSS, and shows the main window."),
    ("MainLauncher.java", "Classpath-friendly launcher for JavaFX."),
    ("DashboardController.java", "Large all-in-one dashboard controller for module screens, navigation, forms, tables, actions, and refresh logic."),
    ("LoginController.java", "Login form controller and transition into the dashboard."),
    ("AdmissionDAO.java", "Admission insert, discharge, active admission lookup, and row mapping."),
    ("AppointmentDAO.java", "Appointment save/update/status/list/search by patient or doctor/date."),
    ("BedDAO.java", "Bed save/status/list/count operations."),
    ("BillingDAO.java", "Bill save/payment/list/revenue operations."),
    ("DBConnection.java", "SQLite connection, schema creation, mock seed data, and last inserted row ID helper."),
    ("DoctorDAO.java", "Doctor/user transactional persistence and doctor queries."),
    ("EmergencyDAO.java", "Emergency queue persistence and priority-sorted retrieval."),
    ("EMRDAO.java", "Medical record persistence and history retrieval."),
    ("HRDAO.java", "Staff shift persistence and listing."),
    ("InventoryDAO.java", "Inventory item and purchase-order persistence."),
    ("LabDAO.java", "Lab request persistence and result updates."),
    ("MedicineDAO.java", "Medicine catalog CRUD, search, and stock updates."),
    ("NotificationDAO.java", "Notification save/read/list/count operations."),
    ("PatientDAO.java", "Patient ID generation, CRUD, soft delete, active list, and search."),
    ("PrescriptionDAO.java", "Prescription header/items transactional save, status update, and retrieval."),
    ("WardDAO.java", "Ward save and list operations."),
    ("AppointmentStatus.java", "Appointment lifecycle enum."),
    ("BedStatus.java", "Bed availability enum."),
    ("PaymentStatus.java", "Payment state enum."),
    ("Role.java", "User role enum."),
    ("TriageLevel.java", "Emergency severity enum with numeric queue priority."),
    ("WardType.java", "Ward category enum."),
    ("Admission.java", "Admission data model."),
    ("Appointment.java", "Appointment data model."),
    ("Bed.java", "Bed data model."),
    ("Bill.java", "Billing data model."),
    ("Doctor.java", "Doctor data model."),
    ("EmergencyPatient.java", "Emergency queue data model."),
    ("InventoryItem.java", "Inventory item data model."),
    ("LabRequest.java", "Lab request data model."),
    ("MedicalRecord.java", "EMR data model."),
    ("Medicine.java", "Medicine data model."),
    ("Notification.java", "Notification data model."),
    ("Patient.java", "Patient data model."),
    ("Prescription.java", "Prescription header model."),
    ("PrescriptionItem.java", "Prescription line item model."),
    ("PurchaseOrder.java", "Restock order data model."),
    ("StaffShift.java", "Staff schedule data model."),
    ("User.java", "Authenticated user/account data model."),
    ("Ward.java", "Ward data model."),
    ("AppointmentService.java", "Appointment validation, booking, conflict prevention, and status changes."),
    ("AuthService.java", "Login, current-user session state, and logout."),
    ("BillingService.java", "Bill calculation, invoice creation, payment processing, and revenue data."),
    ("DoctorService.java", "Doctor profile/account validation and service methods."),
    ("EmergencyService.java", "Emergency triage validation, priority assignment, bed assignment, and resolution."),
    ("EMRService.java", "Medical record validation and retrieval service."),
    ("HRService.java", "Staff shift validation and management service."),
    ("InventoryService.java", "Inventory validation, stock adjustment, purchase order creation, and low-stock alerts."),
    ("LabService.java", "Lab request/result validation and retrieval service."),
    ("PatientService.java", "Patient validation, registration, update, search, and soft delete service."),
    ("PharmacyService.java", "Medicine management, prescription creation, stock checking, and dispensing service."),
    ("ReportsService.java", "Dashboard statistics and report aggregates."),
    ("WardService.java", "Admission/discharge workflow and bed/patient status coordination."),
    ("AlertHelper.java", "Themed information, warning, error, and confirmation dialogs."),
    ("HashHelper.java", "SHA-256 password hashing."),
    ("Validator.java", "Shared validation helper methods."),
    ("LoginView.fxml", "Login screen layout."),
    ("DashboardView.fxml", "Dashboard shell and sidebar/content layout."),
    ("theme.css", "Simple white/blue/gray JavaFX stylesheet."),
    ("DBConnectionTest.java", "JUnit test that verifies database connection, schema, and seed data."),
]
add_table(doc, ["File", "Description"], all_files, [2800, 6560])

doc.add_heading("19. How to Run the Project", level=1)
doc.add_paragraph("Eclipse-native run path:")
add_numbered(doc, [
    "Open Eclipse.",
    "Select File > Import > Existing Projects into Workspace.",
    "Choose the project folder: F:\\Hospital Management System OOP Project.",
    "Import HospitalManagementSystem.",
    "Run HospitalManagementSystem.launch or run com.hospital.MainLauncher.",
])
doc.add_paragraph("Gradle run path, if desired:")
add_numbered(doc, [
    "Ensure Java 17 is installed.",
    "Run gradle run from the project root, or use the Gradle task inside an IDE.",
    "Gradle will use build.gradle to resolve JavaFX and SQLite dependencies.",
])
doc.add_paragraph("Default sample credentials shown on the login screen:")
add_bullets(doc, ["admin / admin123", "dr_smith / doctor123"])

doc.add_heading("20. Notes on Recent Fixes and Design Changes", level=1)
add_bullets(doc, [
    "The project was made Eclipse-native by adding Eclipse project metadata and copying JavaFX/SQLite jars into lib.",
    "The SQLite generated-key problem was fixed by avoiding unsupported Statement.RETURN_GENERATED_KEYS calls and using SQLite's last_insert_rowid() through DBConnection.getLastInsertId(conn).",
    "The visual theme was simplified from a dark/purple style to a white interface with black text, blue active/primary buttons, and gray inactive navigation buttons.",
    "The compiled bin/main resources were refreshed so Eclipse runs with the updated CSS immediately.",
])

doc.add_heading("21. Maintenance Guidance", level=1)
add_bullets(doc, [
    "Keep model classes simple and avoid putting SQL or UI code in them.",
    "Put new business rules in services rather than controllers or DAOs.",
    "Put direct SQL in DAOs and keep query/result mapping localized there.",
    "If a new database table is added, update DBConnection.createTables, add a model, DAO, service methods, and any UI screens.",
    "For SQLite auto-increment inserts, continue using DBConnection.getLastInsertId(conn) after executeUpdate.",
    "When changing CSS, edit src/main/resources/css/theme.css and copy resources into bin/main or refresh/build in Eclipse.",
    "Avoid editing generated folders such as .gradle and build unless troubleshooting build output.",
])

doc.add_heading("22. Quick Project Map", level=1)
doc.add_paragraph(
    "Screen -> Controller -> Service -> DAO -> SQLite -> DAO maps rows -> Service returns result -> Controller refreshes screen."
)
add_table(doc, ["Example", "Path"], [
    ("Register patient", "Patient dialog -> PatientService.registerPatient -> PatientDAO.save -> patients table -> table refresh."),
    ("Book appointment", "Appointment form -> AppointmentService.bookAppointment -> AppointmentDAO.getByDoctorIdAndDate conflict check -> AppointmentDAO.save."),
    ("Dispense prescription", "Pharmacy screen -> PharmacyService.dispensePrescription -> stock checks -> MedicineDAO.updateStock -> PrescriptionDAO.updateStatus."),
    ("Admit patient", "Ward screen -> WardService.admitPatient -> BedDAO.updateStatus + PatientDAO.update + AdmissionDAO.save."),
    ("Resolve ER case", "ER screen -> EmergencyService.resolveEmergency -> EmergencyDAO.updateStatus."),
], [2600, 6760])

for section in doc.sections:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Hospital Management System Project Report")

doc.save(OUT)
print(OUT)
